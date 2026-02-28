import io
import docx
from pypdf import PdfReader
from openai import OpenAI
import streamlit as st

def extract_text(uploaded_file, file_type):
    """Safely extracts raw text from uploaded PDF or DOCX files."""
    text = ""
    try:
        if file_type == "pdf":
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                # Add a fallback for empty pages
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        st.error(f"Error reading your file. Please ensure it is a valid PDF or DOCX. Details: {e}")
    
    return text

def generate_tailored_resume(job_title, job_desc, original_resume):
    """Calls OpenAI to optimize the resume based on the job description."""
    # Pulls the API key directly from Streamlit's secure vault
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    
    # The 'Secret Sauce' Prompt
    system_prompt = (
        "You are an elite Executive Recruiter and ATS (Applicant Tracking System) Expert. "
        "Your task is to rewrite the user's resume to perfectly align with the provided job description.\n\n"
        
        "CORE PHILOSOPHY:\n"
        "1. MAXIMUM RETENTION: You MUST use as much of the original resume's content as possible. "
        "Preserve real company names, dates, and core roles. Do not delete valid experience unless it is 100% irrelevant.\n"
        "2. RELEVANCE & SYNERGY: Rephrase existing bullet points to use the exact terminology found in the Job Description. "
        "If the JD asks for 'Data Visualization' and the resume says 'Made charts,' change it to 'Designed high-impact Data Visualizations.'\n"
        "3. STRATEGIC HYPOTHETICALS: If a critical skill from the JD is missing, add a 'truth-adjacent' bullet point "
        "to an existing role that logically could have included that skill. It must look 100% authentic to the candidate's career path.\n"
        "4. BELIEVABILITY: Maintain a professional, grounded tone. Avoid 'fluff' words like 'passionate' or 'motivated.' "
        "Use concrete results and action verbs.\n\n"
        
        "CRITICAL RULES:\n"
        "1. Output ONLY the resume text. No conversational filler.\n"
        "2. Structure: Name/Contact, Professional Summary, Core Skills, Experience, Education.\n"
        "3. Quantify results with metrics (e.g., 'Optimized SQL queries, reducing report latency by 20%').\n"
        "4. FORMATTING: Use '###' for headers and '-' for bullet points."
    )
    
    user_prompt = f"TARGET ROLE: {job_title}\n\nJOB DESCRIPTION:\n{job_desc}\n\nORIGINAL RESUME:\n{original_resume}"
    
    # Call gpt-4o-mini
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.7 # 0.7 gives a good balance of creativity and factual grounding
    )
    
    return response.choices[0].message.content

def create_word_doc(generated_text):
    """
    Parses the AI's Markdown-style text and converts it into a 
    beautifully formatted, downloadable Word document.
    """
    doc = docx.Document()
    
    # Parse the text line by line to apply native Word formatting
    for line in generated_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # If the AI generated a header
        if line.startswith('###') or line.startswith('**'):
            clean_text = line.replace('###', '').replace('**', '').strip()
            doc.add_heading(clean_text, level=2)
            
        # If the AI generated a bullet point
        elif line.startswith('-') or line.startswith('*'):
            clean_text = line.replace('-', '', 1).replace('*', '', 1).strip()
            doc.add_paragraph(clean_text, style='List Bullet')
            
        # Normal paragraph text (like the summary)
        else:
            doc.add_paragraph(line)
            
    # Save the document to an in-memory byte stream (no files saved to disk!)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio
